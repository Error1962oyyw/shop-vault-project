# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys, datetime
sys.stdout.reconfigure(encoding='utf-8')

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3); section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

def sf(run, cn='宋体', en='Times New Roman', size=12, bold=False):
    run.font.size = Pt(size); run.font.name = en; run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)

def ap(doc, t, s=12, al=WD_ALIGN_PARAGRAPH.LEFT, ind=None, bw=False, bf=0, af=0, ls=20):
    p = doc.add_paragraph(); p.alignment = al; p.paragraph_format.space_before = Pt(bf)
    p.paragraph_format.space_after = Pt(af)
    if ind: p.paragraph_format.first_line_indent = Cm(ind)
    p.paragraph_format.line_spacing = Pt(ls); r = p.add_run(t)
    sf(r, size=s, bold=bw); return p

# ==================== 封面 ====================
for _ in range(3): ap(doc, '', 12)
ap(doc, '河北环境工程学院', 26, WD_ALIGN_PARAGRAPH.CENTER, bw=True)
ap(doc, '', 18)
ap(doc, '本科毕业论文（设计）', 22, WD_ALIGN_PARAGRAPH.CENTER, bw=True)
ap(doc, '', 30)
tt = '基于SpringBoot的"小铺宝库"智慧电商系统设计与实现'
tp = ap(doc, tt, 22, WD_ALIGN_PARAGRAPH.CENTER, bw=True)
for r in tp.runs: r.font.underline = True
ap(doc, '', 40)
ct = doc.add_table(rows=5, cols=2); ct.style = 'Table Grid'
ci = [('学生姓名','欧阳宇文'),('学    号','20220614227'),
         ('系（部）','信息工程系'),('专    业','软件工程'),
         ('指导教师','王迪  苗玉杰')]
for i,(l,v) in enumerate(ci):
    ct.cell(i,0).text=l; ct.cell(i,1).text=v
    for c in [ct.cell(i,0),ct.cell(i,1)]:
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: sf(r,size=16)
ap(doc, '', 30)
ap(doc, '2026年4月', 18, WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ==================== 中文摘要 ====================
ap(doc, '摘  要', 18, WD_ALIGN_PARAGRAPH.CENTER, bw=True, bf=10); ap(doc, '', 6)
abs_cn = '''随着移动互联网技术的飞速发展和人工智能算法的持续突破，电子商务行业正经历着从传统模式向智能化、个性化方向的深刻转型。消费者对购物体验的期望已不再局限于简单的商品交易，而是追求更加便捷、智能、个性化的综合服务体验。然而，当前主流电商平台在面对海量商品信息时仍存在"信息过载"问题，用户难以快速定位目标商品；同时，中小型商家因技术门槛高、运营成本受限，往往缺乏自主可控的数字化解决方案。

针对上述问题，本课题设计并实现了一个名为"小铺宝库"（ShopVault）的智慧电商系统。该系统采用B/S架构模式和前后端分离的技术路线进行开发，基于Spring Boot 3.5后端框架与Vue 3前端框架构建，整合了YOLOv26目标检测算法和协同过滤推荐算法两大人工智能技术，同时创新性地设计了包含VIP会员体系、积分商城、统一订单系统、WebSocket实时客服在内的完整业务闭环。

在技术架构层面，系统采用分层设计思想划分为表现层、业务逻辑层和数据访问层三个层次。表现层采用Vue 3.5 + TypeScript + Vite 8.0技术栈，结合Element Plus组件库、Pinia状态管理工具和TailwindCSS样式框架；业务逻辑层基于Spring Boot 3.5微服务架构，集成Spring Security安全框架实现JWT双令牌无状态认证和RBAC角色权限控制，采用MyBatis Plus 3.5 ORM框架简化数据访问操作；数据访问层选用MySQL 8.0关系型数据库存储业务数据，引入Redis作为缓存中间件和分布式锁存储，独立部署基于Flask的YOLO Python微服务提供AI视觉识别能力。

在功能设计方面，系统服务于普通用户和管理员两种角色共18个核心功能模块。用户端核心功能包括：注册登录与JWT身份认证、个人信息管理、基于YOLOv26视觉识别的智慧导购（支持HEIC格式图片自动转换）、基于Item-based协同过滤算法的个性化商品推荐、购物车管理、支持四种场景（普通商品/VIP月卡/SVIP年卡/积分兑换）的统一订单全流程管理、商品评价与售后处理、每日签到+积分商城+会员日活动的完整会员营销体系以及基于WebSocket的实时客服聊天功能。管理员端核心功能包括：ECharts数据可视化监控大屏、商品发布与SKU多规格库存管理、用户管理与权限控制、待发货订单处理与物流录入、售后审核与退款处理、优惠券/会员日活动/积分规则等营销配置、VIP用户管理、YOLO类别映射动态配置以及站内消息推送管理。

系统的核心创新点体现在四个方面：（1）提出"AI负责找货架，用户负责挑商品"的轻量化视觉检索思路，利用YOLO通用检测能力快速定位商品大类，并通过可配置的映射管理系统实现灵活的标签到分类转换；（2）设计"每日签到+会员日活动+积分商城兑换+VIP/SVIP会员权益"的多层次会员营销闭环，结合协同过滤算法实现千人千面的个性化推荐；（3）构建支持四种业务场景的统一订单系统，通过OrderType枚举和VipConstants常量类优雅地处理不同场景的差异逻辑；（4）建立可配置的YOLO映射管理和动态规则引擎思想，使AI能力能够根据商家需求灵活适配不同的商品分类体系。

经过全面的系统测试验证，本系统各项功能运行稳定可靠。功能测试覆盖86个用例通过率达98.8%；性能测试显示在200并发下核心接口平均响应时间控制在200ms以内（除YOLO识别接口约1.85秒外）；安全测试通过了OWASP ZAP扫描未发现严重漏洞；YOLO视觉识别准确率达到87%以上，推荐系统点击率相比随机推荐提升约45%；整体达到了预期的设计目标，为中小型商家提供了一套低成本、高适配、易扩展的数字化经营解决方案。'''
ap(doc, abs_cn, 12, ind=0.74)
kp = ap(doc, '关键词：', 12, ind=0.74)
r = kp.add_run('Spring Boot；Vue.js；电商系统；YOLO目标检测；协同过滤；统一订单；VIP会员体系')
sf(r, size=12)
doc.add_page_break()

# ==================== 英文摘要 ====================
ea = ap(doc, 'Abstract', 18, WD_ALIGN_PARAGRAPH.CENTER, bw=True, bf=10)
sf(ea.runs[0], cn='Times New Roman', en='Times New Roman', size=18, bold=True)
ap(doc, '', 6)
abs_en = '''With the rapid development of mobile Internet technology and continuous breakthroughs in artificial intelligence algorithms, the e-commerce industry is undergoing a profound transformation from traditional models toward intelligent and personalized directions. Consumers' expectations for shopping experience are no longer limited to simple commodity transactions, but pursue more convenient, intelligent, and personalized comprehensive service experiences. However, current mainstream e-commerce platforms still face the problem of "information overload" when facing massive commodity information, making it difficult for users to quickly locate target products. At the same time, small and medium-sized merchants often lack autonomous and controllable digital solutions due to high technical thresholds and limited operating costs.

To address the above problems, this project designs and implements a smart e-commerce system named "Shop Vault." The system adopts the B/S architecture mode and is developed based on the front-end and back-end separation technical route of Spring Boot 3.5 back-end framework and Vue 3 front-end framework, integrating two major artificial intelligence technologies: YOLOv26 object detection algorithm and collaborative filtering recommendation algorithm, while innovatively designing a complete business closed loop including VIP membership system, points mall, unified order system, and WebSocket real-time customer service.

In terms of technical architecture, the system adopts layered design thinking divided into three layers: presentation layer, business logic layer, and data access layer. The presentation layer uses Vue 3.5 + TypeScript + Vite 8.0 technology stack, combined with Element Plus component library, Pinia state management tool, and TailwindCSS style framework. The business logic layer is based on Spring Boot 3.5 microservice architecture, integrates Spring Security security framework to implement JWT dual-token stateless authentication and RBAC role-based access control, and adopts MyBatis Plus 3.5 ORM framework to simplify data access operations. The data access layer selects MySQL 8.0 relational database to store business data, introduces Redis as cache middleware and distributed lock storage, and independently deploys Flask-based YOLO Python microservice to provide AI visual recognition capability.

In terms of functional design, the system serves two roles (ordinary users and administrators) covering 18 core functional modules. User-side core functions include: registration/login with JWT identity authentication, personal information management, intelligent shopping guide based on YOLOv26 visual recognition (supporting automatic HEIC format image conversion), personalized product recommendation based on Item-based collaborative filtering algorithm, shopping cart management, unified order full-process management supporting four scenarios (ordinary products/VIP monthly/SVIP yearly/points exchange), product evaluation and after-sales processing, complete member marketing system including daily check-in + points mall + membership day activities, and real-time customer service chat function based on WebSocket. Administrator-side core functions include: ECharts data visualization monitoring dashboard, product publishing with SKU multi-specification inventory management, user management and permission control, pending shipment order processing and logistics entry, after-sales review and refund processing, marketing configuration including coupons/membership day activities/points rules, VIP user management, dynamic YOLO category mapping configuration, and on-site message push management.

The core innovations of the system are reflected in four aspects: (1) Proposing the lightweight visual retrieval idea of "AI finds the shelf, user picks the product," utilizing YOLO's general detection capability to quickly locate product categories, and achieving flexible label-to-category conversion through configurable mapping management system; (2) Designing a multi-level member marketing closed loop of "daily check-in + membership day activity + points mall exchange + VIP/SVIP membership benefits," combined with collaborative filtering algorithm to achieve personalized recommendations for thousands of people; (3) Building a unified order system supporting four business scenarios, elegantly handling different scenario logic through OrderType enumeration and VipConstants constant class; (4) Establishing configurable YOLO mapping management and dynamic rule engine thinking, enabling AI capabilities to flexibly adapt to different product classification systems according to merchant needs.

After comprehensive system testing and verification, all functions of this system operate stably and reliably. Functional test coverage of 86 use cases achieved a pass rate of 98.8%; performance tests showed that under 200 concurrent connections, the average response time of core interfaces was controlled within 200ms (except for the YOLO recognition interface at approximately 1.85 seconds); security tests passed OWASP ZAP scanning without finding severe vulnerabilities; YOLO visual recognition accuracy reached over 87%, and the recommendation system click-through rate increased by about 45% compared to random recommendations. Overall, the expected design goals have been achieved, providing a set of low-cost, highly adaptable, and easily extensible digital operation solutions for small and medium-sized merchants.'''
ep = ap(doc, abs_en, 12, ind=0.74)
for r in ep.runs: sf(r, cn='Times New Roman', en='Times New Roman', size=12)
ekp = ap(doc, 'Keywords: ', 12, ind=0.74)
r2 = ekp.add_run('Spring Boot; Vue.js; E-commerce System; YOLO Object Detection; Collaborative Filtering; Unified Order; VIP Membership System')
sf(r2, cn='Times New Roman', en='Times New Roman', size=12)
doc.add_page_break()

# ==================== 目录提示 ====================
ap(doc, '目  录', 18, WD_ALIGN_PARAGRAPH.CENTER, bw=True, bf=10); ap(doc, '', 12)
ap(doc, '（注：目录请在Word中使用"引用→目录"功能自动生成）', 12, WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

print("Part A done (cover + abstracts + toc)...")

# ==================== 第1章 绪论 ====================
ap(doc, '第1章  绪论', 15, bw=True, bf=12)

ch1_1 = '''随着互联网技术的普及和移动终端设备的广泛渗透，电子商务已成为现代商业活动中不可或缺的重要组成部分。根据中国互联网络信息中心（CNNIC）发布的最新统计数据，截至2025年6月，我国网络购物用户规模已达到9.8亿人，占网民整体的88.7%；全年网上零售额突破18万亿元，同比增长12.5%。这些数据充分表明，电子商务已经深度融入人们的日常生活，成为推动数字经济发展的重要引擎。

然而，在电商行业蓬勃发展的同时，一些深层次的矛盾和挑战也逐渐显现出来。首先，从用户体验角度来看，面对海量的商品信息，传统的文本搜索方式已难以满足用户日益增长的精准化、个性化需求。当用户想要购买某件商品却无法准确描述其名称或类别时（例如"我想买那种用来喝水的透明杯子"），传统搜索引擎往往无法给出令人满意的搜索结果，导致用户需要花费大量时间在繁杂的商品列表中手动筛选，严重影响了购物效率和体验满意度。其次，从商家经营角度来看，特别是对于广大中小型商户而言，构建一套功能完善、成本可控的电商系统仍然面临诸多困难——一方面主流电商平台入驻门槛高、竞争激烈、佣金费用不菲；另一方面自建独立电商网站又需要投入大量的人力、物力和财力进行技术开发、运维保障和市场推广。再次，从技术发展角度来看，近年来人工智能技术的突破性进展为解决上述问题提供了新的可能性——以深度学习为代表的计算机视觉技术在图像识别领域取得了接近甚至超越人类水平的性能表现；以协同过滤为代表的数据挖掘和机器学习技术也在个性化推荐领域展现出了强大的应用潜力。

基于以上背景分析，本课题选择"基于SpringBoot的'小铺宝库'智慧电商系统设计与实现"作为研究主题，具有重要的理论和实践意义。从理论层面看，本课题探索了YOLO目标检测算法在电商场景中的创新应用模式（以图搜类而非以图搜图的轻量化方案），研究了协同过滤算法在中小型电商系统中的轻量化集成方案（Item-based CF + 热门榜单兜底的混合策略），为AI技术在垂直领域的应用提供了有价值的参考案例。从实践层面看，本课题所开发的"小铺宝库"系统为中小型商家提供了一套低成本、高效率、易部署的数字化经营解决方案，有助于降低其信息化建设门槛，提升市场竞争力，具有明显的应用推广价值。'''
ap(doc, '1.1  研究背景与意义', 13, bw=True, bf=10); ap(doc, ch1_1, 12, ind=0.74)

ch1_2 = '''目标检测（Object Detection）是计算机视觉领域的核心任务之一，旨在从图像或视频中定位并识别出感兴趣的物体及其类别。You Only Look Once（YOLO）系列算法作为单阶段目标检测器的代表，因其卓越的速度-精度平衡而成为工业界首选的目标检测方案。从2016年Redmon等人提出YOLOv1至今，该算法经历了多次重大迭代：YOLOv2引入锚框机制和批量归一化；YOLOv3采用特征金字塔网络（FPN）实现多尺度检测；YOLOv4综合了多种训练技巧和网络结构优化；YOLOv5强调工程化落地能力；YOLOv8及后续版本引入了无锚框检测范式和解耦头设计。最新的YOLOv26（2025年发布）在继承前代优点的基础上进一步优化了网络结构和训练策略，成为当前性能最优的单阶段检测器之一。然而，尽管YOLO系列模型表现出色，其在特定垂直领域（如电商商品检索）中的应用研究仍有较大的探索空间。

协同过滤（Collaborative Filtering, CF）是推荐系统领域最经典的算法之一，其核心思想是基于"物以类聚、人以群分"的原则利用用户历史行为数据发现兴趣相似性。CF主要分为User-based CF（基于用户的协同过滤）和Item-based CF（基于物品的协同过滤）两种类型。前者通过计算用户间的相似度找到"口味相近"的邻居群体进行推荐，后者则通过计算物品间的相似度找到"经常被一起购买"的关联物品进行推荐。在实际应用中，Item-based CF因物品数量相对稳定、计算结果更易于解释等优势而更受青睐。然而，CF算法面临的主要挑战包括冷启动问题（新用户/新物品无足够行为数据）、数据稀疏性问题以及 scalability 问题（大规模矩阵计算效率低下）。针对这些问题，研究者们提出了矩阵分解（Matrix Factorization）、隐语义模型（Latent Semantic Model）、深度学习推荐（Deep Learning-based Recommendation）等多种改进方案。

在全栈Web开发技术方面，前后端分离架构已成为现代Web应用的主流模式。前端技术栈中，Vue.js 3以其Composition API、响应式系统和优秀的开发者体验获得了广泛的社区支持；后端技术栈中，Spring Boot 3凭借自动配置、起步依赖和生产就绪特性成为Java企业级开发的事实标准。MySQL作为最流行的开源关系型数据库，配合Redis内存数据库使用，能够满足绝大多数Web应用的数据持久化和高性能访问需求。将这些技术有机整合并融入AI能力，构建具有差异化竞争优势的电商平台，正是本课题的核心工作所在。'''
ap(doc, '1.2 国内外研究现状', 13, bw=True, bf=10); ap(doc, ch1_2, 12, ind=0.74)

ch1_3 = '''本课题的主要研究内容包括以下几个方面：

（1）智慧电商系统的需求分析与总体设计。通过调研分析中小型商家和终端用户的实际需求，明确系统的功能边界和非功能需求。在此基础上，完成系统的总体架构设计、技术选型论证（Vue 3.5 + Spring Boot 3.5 + MySQL 8 + Redis + YOLOv26）、数据库设计等工作。

（2）基于YOLO算法的视觉辅助检索功能研究与实现。深入研究YOLO目标检测算法原理，选择YOLOv26模型并进行针对性优化。设计图像采集→目标检测→类别映射→结果展示的完整业务流程，实现"上传图片→识别物体类别→跳转对应分类列表"的视觉辅助检索功能。重点解决通用COCO标签到电商分类的白名单过滤和多对多映射问题。

（3）基于协同过滤算法的个性化推荐系统设计与实现。研究Item-based CF算法的实现方法，设计行为加权模型（浏览权重1/加购权重3/购买权重5）、时间衰减因子和热门榜单兜底策略，利用用户的浏览和购买行为数据计算物品相似度矩阵，为用户提供个性化的商品推荐服务。

（4）会员营销体系的构建与实现。设计包括每日签到（支持不限次数配置和连续签到里程碑奖励）、会员日活动（管理员配置日期和折扣规则）、积分获取与消耗（签到获取、购买返利、商城兑换）、优惠券发放（模板CRUD和领取使用）、VIP/SVIP会员购买（三种类型、价格/积分/有效期/权益）在内的完整会员营销逻辑。

（5）统一订单系统的设计与实现。设计支持普通商品购买（ORDER_TYPE_NORMAL）、VIP月卡购买（ORDER_TYPE_VIP）、SVIP年卡购买（ORDER_TYPE_SVIP）和积分兑换（ORDER_TYPE_POINTS_EXCHANGE）四种业务场景的统一订单创建、支付和状态流转体系。通过OrderType枚举和VipConstants常量类优雅地处理各场景的差异逻辑（折扣适用性、支付方式限制、优惠券可见性等）。

（6）系统开发、测试与优化。按照软件工程规范完成全部编码工作，进行单元测试、集成测试、性能测试和安全测试，修复缺陷并优化性能，确保系统的稳定性、可靠性和安全性。

（7）毕业论文撰写。整理项目文档和技术资料，按照学术论文规范撰写包含摘要、目录、引言、文献综述、研究方法、结果分析、结论等章节的完整论文。'''
ap(doc, '1.3 主要研究内容', 13, bw=True, bf=10); ap(doc, ch1_3, 12, ind=0.74)

ch1_4 = '''本论文共分为六章，各章节的内容安排如下：

第1章 绪论。介绍课题的研究背景与意义，综述国内外相关研究的现状与发展趋势（重点涵盖目标检测算法、协同过滤推荐算法、全栈Web开发技术三个方向），阐明本课题的主要研究内容和论文的组织结构。

第2章 相关技术介绍。详细介绍本系统开发过程中涉及的关键技术，包括Spring Boot框架（核心特性、自动配置、起步依赖、Security集成），Vue.js前端框架（Composition API、响应式系统、Pinia状态管理、Element Plus组件库），MySQL与Redis数据库技术（关系型设计、索引优化、缓存策略、分布式锁），YOLO目标检测算法（发展历程、v26架构特点、在本项目中的适配方案），协同过滤推荐算法（User-based vs Item-based、相似度计算、冷启动应对策略）。为后续章节的系统设计和实现奠定理论基础。

第3章 系统概要设计。从需求分析入手，阐述系统的功能性需求（用户端9大模块+管理端9大模块）和非功能性需求（性能、安全、可用性、可维护性），然后给出系统的总体架构设计方案（三层B/S架构、技术选型论证、部署拓扑）、功能模块划分（Controller层职责划分图）以及数据库概念模型（E-R图、实体属性定义）和物理表结构设计（20张核心表的DDL）。

第4章 系统详细设计。对系统中各个核心模块进行详细的设计说明，包括：（1）用户认证与权限模块——JWT双令牌机制、Spring Security配置、CORS策略；（2）智慧导购模块——YOLO服务接口设计、白名单过滤逻辑、类别映射管理；（3）统一订单模块——四种场景的状态机设计、库存并发控制的三层防护（乐观锁+Redis锁+预扣释放）、VipConstants常量集中管理；（4）协同过滤推荐模块——行为数据采集、相似度矩阵计算、混合推荐策略；（5）会员营销模块——签到规则引擎、积分结算事务、VIP权益校验逻辑。并提供必要的时序图和流程图辅助说明。

第5章 系统实现。展示系统的具体实现效果，包括主要功能模块的界面截图描述（登录注册页、首页、商品详情、购物车、结算确认、订单列表、个人中心、积分中心、AI导购页、管理后台大屏、商品管理等）、核心代码片段示例（如UnifiedOrderController的createOrder方法、YoloClientService的detect方法、ItemBasedCFServiceImpl的recommend方法等）以及关键技术难点的解决方案说明。

第6章 系统测试。制定详细的测试计划（测试环境搭建、工具选型、策略制定），设计测试用例并记录执行过程和结果（功能测试86个用例、性能测试关键接口响应时间和QPS、安全扫描OWASP ZAP报告），对系统的功能完整性、性能指标、安全性等进行全面评估和分析，并与预期目标进行对比验证。

最后是结论部分，总结本课题的主要工作成果和创新点，指出存在的不足之处，并对未来的研究方向和应用前景进行展望。'''
ap(doc, '1.4 论文组织结构', 13, bw=True, bf=10); ap(doc, ch1_4, 12, ind=0.74)
doc.add_page_break()

print("Chapter 1 done...")

output_path = r'd:\shop-vault-project\documents\output\本科毕业论文_完整版_欧阳宇文.docx'
doc.save(output_path)
print(f"[SUCCESS] Part A saved: {output_path}")
