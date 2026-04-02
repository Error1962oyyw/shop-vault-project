# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = Document(r'd:\shop-vault-project\documents\output\本科毕业论文_完整版_欧阳宇文.docx')

def sf(run, cn='宋体', en='Times New Roman', size=12, bold=False):
    run.font.size = Pt(size); run.font.name = en; run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)

def ap(doc, t, s=12, al=WD_ALIGN_PARAGRAPH.LEFT, ind=None, bw=False, bf=0, af=0, ls=20):
    p = doc.add_paragraph(); p.alignment = al; p.paragraph_format.space_before = Pt(bf)
    p.paragraph_format.space_after = Pt(af)
    if ind: p.paragraph_format.first_line_indent = Cm(ind)
    p.paragraph_format.line_spacing = Pt(ls); r = p.add_run(t)
    sf(r, size=s, bold=bw); return p

# ==================== 第2章 相关技术介绍 ====================
ap(doc, '第2章  相关技术介绍', 15, bw=True, bf=12)
intro2 = '''本章将对本系统开发过程中所使用的关键技术进行详细介绍，包括后端开发框架Spring Boot、前端开发框架Vue.js、数据库技术MySQL与Redis、目标检测算法YOLOv26以及推荐算法协同过滤等。通过本章的学习，读者可以了解各项技术的基本原理、特点以及在本次系统开发中的具体应用方式。'''
ap(doc, intro2, 12, ind=0.74)

# 2.1 Spring Boot框架
ap(doc, '2.1 Spring Boot框架', 13, bw=True, bf=10)
sb_content = '''Spring Boot是由Pivotal团队（现归VMware）于2014年推出的全新开源Java框架，其设计目标是简化Spring应用的初始搭建和开发过程。该框架遵循"约定优于配置"（Convention over Configuration）的理念，通过自动配置机制和起步依赖（Starter Dependencies）管理，使开发者能够快速构建生产级别的Spring应用程序。

2.1.1 Spring Boot的核心特性

（1）自动配置（Auto-Configuration）：Spring Boot能够根据项目中添加的依赖jar包自动配置Spring应用所需的Bean组件。例如，当classpath下存在spring-boot-starter-web依赖时，Spring Boot会自动配置DispatcherServlet、内嵌Tomcat服务器、HTTP消息转换器等Web开发所需的组件。开发者只需在application.yml中进行少量自定义配置即可覆盖默认行为。

（2）起步依赖（Starter Dependencies）：Spring Boot提供了一系列预定义的依赖集合（称为Starter），每个Starter包含了实现特定功能所需的一组依赖及其经过验证的兼容版本。本系统使用的核心Starter包括：spring-boot-starter-web（Web MVC开发）、spring-boot-starter-security（安全认证集成）、spring-boot-starter-data-redis（Redis缓存集成）、mybatis-plus-spring-boot3-starter（MyBatis Plus ORM）、spring-boot-starter-websocket（WebSocket实时通讯）以及spring-boot-starter-mail（邮件服务）。

（3）内嵌服务器（Embedded Server）：Spring Boot支持内嵌Tomcat、Jetty或Undertow等Servlet容器。应用可以打包为可执行的JAR文件直接通过java -jar命令运行，无需外部安装和配置独立的Web服务器，极大简化了部署流程。本系统在生产环境中采用Nginx反向代理+JAR包部署方案。

（4）生产就绪特性（Production-Ready Features）：Spring Boot Actuator模块提供了丰富的运维端点（/health、/metrics、/info、/env等），便于监控应用的健康状态、性能指标和环境信息。本系统在开发环境开启Actuator以辅助调试，在生产环境中禁用了敏感端点以确保安全。

2.1.2 本系统中的Spring Boot应用架构

本系统后端基于Spring Boot 3.5.12版本构建，项目坐标为com.TsukasaChan.ShopVault，采用标准的分层架构组织代码：

Controller层（控制器层）：负责接收HTTP请求、参数校验、调用Service层并返回JSON响应。本系统共包含29个Controller类，按职责划分为以下子包：
- controller/system/：AuthController（认证注册登录）、UserController（用户信息）、ChatController（WebSocket聊天）
- controller/product/：ProductController（商品CRUD）、SpecController（规格值）、CategoryController（分类）
- controller/order/：CartController（购物车）、OrderController（订单操作）、AfterSalesController（售后）、UnifiedOrderController（统一订单）
- controller/marketing/：MarketingController（签到积分）、VipController（VIP会员）、PointsMallController（积分商城）、ActivityController（活动领券）、CouponController（优惠券）、MemberDaysController（会员日）、PointsController（积分记录）、PointsProductsController（积分商品）、PointsRulesController（积分规则）
- controller/admin/：AdminController（用户管理）、AdminOrderController（订单处理）、AdminAfterSalesController（售后审核）、AdminActivityController（活动管理）、AdminMessageController（消息管理）、AdminPointsProductController（积分商品管理）、AdminPointsRuleController（积分规则管理）、AdminVipController（VIP用户）、DashboardController（数据大屏）、UsersController（余额调整）、YoloMappingController（YOLO映射配置）

Service层（服务层）：实现具体的业务逻辑，包括订单状态流转、积分计算规则、推荐算法调用、库存并发控制等。关键Service包括：UserService、ProductService、UnifiedOrderService、ItemBasedCFServiceImpl（协同过滤推荐）、YoloClientService（YOLO服务集成）、VipBenefitService（VIP权益校验）等。

Security层（安全层）：基于Spring Security + JWT实现无状态认证。SecurityConfig配置了STATELESS会话模式、公开路径白名单（/api/auth/**、/api/category/**、/api/product/list等）、JwtAuthenticationFilter过滤器链、@PreAuthorize("hasRole('ADMIN')")管理员权限注解。密码使用BCrypt算法加密存储，Refresh Token存入Redis实现主动失效能力。

Integration层（集成层）：封装对外部服务的调用。YoloClientService通过RestTemplate调用独立部署的YOLO Python Flask服务（端口5000），实现了图片上传→Base64编码→模型推理→结果解析的完整链路，特别支持HEIF/HEIC格式图片的自动转换处理。'''
ap(doc, sb_content, 12, ind=0.74)

# 2.2 Vue.js前端框架
ap(doc, '2.2 Vue.js前端框架', 13, bw=True, bf=10)
vue_content = '''Vue.js是一款用于构建用户界面的渐进式JavaScript框架，由尤雨溪（Evan You）于2014年创建并开源发布。Vue的设计理念是"渐进式"——即可以作为库嵌入现有页面逐步增强，也可以作为框架驱动整个单页面应用（SPA）的开发。Vue 3是该框架的最新主要版本，引入了Composition API（组合式API）这一重要的编程范式变革。

2.2.1 Vue 3的核心新特性

（1）Composition API（组合式API）：Vue 3新增了setup()函数语法，允许开发者按照逻辑关注点而非选项类型（data/methods/computed/watch）来组织代码。配合ref()和reactive()响应式API以及computed()和watch()工具函数，可以实现更灵活的逻辑复用和更好的TypeScript类型推导。本系统全部组件均采用<script setup lang="ts">语法编写。

（2）响应式系统重构：Vue 3使用ES6 Proxy对象替代了Vue 2中的Object.defineProperty来实现数据的响应式监听。新系统能够检测属性的动态添加和删除、支持Map/Set等新型数据结构的响应式追踪，并且性能更优（避免了递归遍历带来的开销）。

（3）Tree-shaking友好：Vue 3的核心库对不使用的功能进行了摇树优化（Tree-shaking），使得打包后的体积显著减小。配合Vite构建工具按需导入Element Plus组件（unplugin-vue-components插件），本系统前端打包体积控制在合理范围内。

（4）Teleport组件：Vue 3正式引入了<Teleport>组件（原为Vue 2的第三方vue-teleport），允许将组件渲染到DOM树的任意位置。本系统中用于将全局的加载遮罩、消息提示框等UI元素传送到<body>层级，避免被父容器的overflow:hidden等CSS属性裁剪。

2.2.2 本系统前端技术栈详解

本系统前端基于Vue 3.5.30版本构建，package.json定义了完整的技术依赖：

核心框架：vue@3.5.30、vue-router@4.6.2（路由管理）、pinia@3.0.0（状态管理替代Vuex）
构建工具：vite@8.0.0（极速HMR开发服务器）、typescript@5.9.0（类型安全）
UI组件库：element-plus@2.13.14（企业级组件）、@element-plus/icons-vue（图标）
HTTP客户端：axios@1.13.0（二次封装含拦截器）
样式方案：tailwindcss@4.2.0（原子化CSS类名）
图表库：echarts@5.5.1（数据可视化大屏）

前端目录结构（src/）：
- api/：14个API模块文件（auth.ts、user.ts、product.ts、cart.ts、order.ts、marketing.ts、vip.ts、points-mall.ts、chat.ts等）
- views/：25+个页面组件（home/、login/、register/、products/、product/、cart/、checkout/、orders/、profile/、points/、points-mall/、member-day/、chat/、ai-search/、favorites/、messages/、forgot-password/以及admin/下的16个管理页面）
- stores/：Pinia状态仓库（user.ts用户认证态、cart.ts购物车、category.ts分类缓存）
- types/：TypeScript类型定义（api.d.ts接口类型、user.d.ts用户实体等）
- utils/：工具函数（request.ts Axios封装、auth.ts令牌管理、format.ts日期格式化）
- router/index.ts：路由配置（hash模式、路由守卫权限控制、懒加载）
- App.vue：根组件（<RouterView>容器）
- main.ts：入口文件（createApp挂载插件）'''
ap(doc, vue_content, 12, ind=0.74)

# 2.3 MySQL与Redis数据库技术
ap(doc, '2.3 MySQL与Redis数据库技术', 13, bw=True, bf=10)
db_content = '''2.3.1 MySQL关系型数据库

MySQL是目前全球最流行的开源关系型数据库管理系统（RDBMS），由Oracle公司下属的MySQL AB部门维护和发展。其之所以广受青睐，主要源于以下优势：（1）开源免费——社区版和企业版均可免费使用；（2）高性能——InnoDB存储引擎支持事务ACID、行级锁和外键约束，查询优化器能够高效执行复杂SQL；（3）高可用——支持主从复制（Master-Slave Replication）、组复制（Group Replication）和集群（InnoDB Cluster）等多种高可用方案；（4）生态丰富——与PHP、Java、Python、Node.js等主流语言均有成熟的驱动程序和ORM框架。

本系统选用MySQL 8.0版本作为主数据库，通过application.yml配置连接参数：
- 连接池：HikariCP（Spring Boot 2.x+默认连接池，性能优于Druid）
- 字符集：utf8mb4（支持完整的Unicode包括emoji）
- 时区：Asia/Shanghai
- SQL日志：开启慢查询日志（long_query_time > 2秒）

数据库共设计了20张核心表，按业务域划分如下：

用户域（4张）：sv_user（用户基本信息，含手机号/邮箱/昵称/头像/密码BCrypt加密/角色/状态）、sv_address（收货地址，最多5个）、sv_message（站内消息）、sv_favorites（收藏夹）

商品域（5张）：sv_product（商品主表，含名称/描述/主图URL/价格/原价/库存/销量/评分均值/分类ID/上架状态/SKU启用标志）、sv_category（商品分类，支持二级树形结构parent_id自引用）、sv_spec（商品规格名称如颜色/尺寸）、sv_sku（具体SKU值，关联商品ID和规格ID组合唯一键）、sv_product_image（商品图片，支持多图）

交易域（7张）：sv_cart_item（购物车项）、sv_order（订单主表，含order_no编号/user_id/总金额/实付金额/status状态/pay_type支付方式/收货人信息/created_at下单时间）、sv_order_item（订单明细，关联订单-商品-SKU-数量-小计）、sv_review（评价，1-5星+图文+审核状态）、sv_after_sale（售后申请，含类型/原因/凭证图/审核状态/处理结果）

营销域（8张）：sv_points_record（积分变动记录，正数获取负数消耗）、sv_coupon（优惠券模板，含面额/门槛/有效期/发行总量）、sv_user_coupon（用户领取的优惠券）、sv_points_product（积分兑换商品）、sv_points_rule（签到规则配置，dailyLimit连续奖励）、sv_activity（会员日活动配置）、svip_membership（VIP/SVIP会员记录，含类型/开始时间/到期时间/等级）、sv_dashboard_stat（大屏统计数据缓存）

2.3.2 Redis内存数据库

Redis（Remote Dictionary Server）是一个开源的高性能键值对NoSQL数据库，由Salvatore Sanfilippo于2009年用C语言编写。其数据结构丰富（String/Hash/List/Set/Sorted Set/Stream等），读写速度极快（单线程模型避免锁竞争），持久化支持RDB/AOF两种机制。

在本系统中Redis承担以下五类职责：

（1）会话与令牌缓存：存储用户的JWT Refresh Token（key: refresh:token:{userId}），设置7天过期时间（TTL）。Access Token存放在前端Pinia Store内存中（1小时有效），过期时携带Refresh Token调用刷新接口换取新的Access Token组合。用户主动退出登录时删除Redis中的Token实现"踢出线"效果。

（2）热点数据缓存：缓存首页轮播图列表（key: banner:list）、热门商品Top20（key: product:hot）、商品分类全量树（key: category:tree）。这些数据变更频率低但访问频率高，缓存后可大幅减轻MySQL读取压力。缓存策略采用Cache-Aside模式（先查Redis，未命中则查DB并回写Redis）。

（3）分布式锁实现：在高并发场景（如秒杀活动、库存扣减）中使用Redis的SETNX命令实现互斥锁。锁的key格式为stock:lock:{productId}，value为请求唯一标识（UUID），过期时间设置为15秒（需大于业务最大耗时）。获取锁成功后才执行库存修改操作，防止超卖。

（4）库存预扣与释放：用户下单时将对应数量的库存从MySQL"可用库存"转移到Redis"锁定库存"（key: stock:locked:{orderId}:{productId}）中。支付成功后同步扣减MySQL库存并清除Redis锁定；若超时未支付（定时任务每分钟检查），则自动释放锁定回MySQL可用库存。

（5）推荐结果缓存：协同过滤算法计算的物品相似度矩阵（N×N维，N为商品总数）计算耗时较长，因此在离线任务中每小时更新一次并存入Redis（key: cf:similarity_matrix），在线推荐阶段仅需读取缓存的相似度矩阵并进行加权求和排序即可返回Top-K结果。'''
ap(doc, db_content, 12, ind=0.74)

# 2.4 YOLO目标检测算法
ap(doc, '2.4 YOLO目标检测算法', 13, bw=True, bf=10)
yolo_content = '''2.4.1 YOLO算法发展历程

You Only Look Once（YOLO）是由Joseph Redmon等人于2016年提出的一种革命性实时目标检测算法。与传统两阶段方法（如R-CNN系列先生成候选区域再分类）不同，YOLO将目标检测转化为单一的回归问题：输入一张图像，网络一次性前向传播即可同时预测所有目标的边界框坐标（x, y, w, h）和类别置信度（class confidence）。这种端到端的范式使得YOLO在速度上具有压倒性优势——初代YOLOv1在COCO数据集上即可达到45 FPS（Frame Per Second）的推理速度，同时保持63.4% mAP@0.5的精度。

从YOLOv1到最新的YOLOv26，该算法经历了多次重大迭代升级：

- YOLOv1（2016）：开创性地将检测问题回归化，Darknet-19骨干网络提取特征，直接在全图上预测边界框。
- YOLOv2（2017）：引入Anchor Box机制（预设9种尺寸锚框）和Batch Normalization，mAP提升至76.8%，速度达67 FPS。
- YOLOv3（2018）：采用Darknet-53骨干网络（借鉴ResNet残差思想），引入FPN（Feature Pyramid Network）实现多尺度检测（三个输出层分别负责小/中/大目标）。
- YOLOv4（2020）：综合了Mosaic增强、CutMix、Label Smoothing、CIoU Loss、CSPDarknet53、SPP空间金字塔池化、PANet路径聚合等大量Trick，COCO mAP@0.5:0.95达到43.5%。
- YOLOv5（2020）：由Ultralytics公司用PyTorch重写，强调工程化和易用性，提供n/s/m/l/x五种规模变体。
- YOLOv8/v9/v10/v11（2023-2024）：引入无锚框（Anchor-Free）解耦头（Decoupled Head）、动态标签分配（Dynamic Label Assignment）等先进设计。
- YOLOv26（2025）：最新版本，进一步优化了网络结构和训练策略，在COCO benchmark上刷新了性能记录。

2.4.2 YOLO在本系统中的应用架构

本系统的视觉导购功能采用了YOLOv26模型（Ultralytics官方实现yolo26x.pt），并将其独立部署为一个Python微服务。整体架构如下：

YOLO服务端（Python Flask）：
- 框架：Flask轻量级Web框架
- 端口：5000
- 核心模型：ultralytics.YOLO("yolo26x.pt")
- API端点：POST /predict
- 输入：multipart/form-data格式的图像文件（支持jpg/png/webp/heic/heif）
- 处理流程：
  (1) 接收上传的图像文件
  (2) 使用PIL库读取图像并根据格式判断是否需要HEIC→JPEG转换（通过pillow_heif插件）
  (3) 将图像转换为RGB模式的numpy数组
  (4) 调用model.predict()执行推理
  (5) 遍历检测结果，去重同类标签（同一物体多次检测只保留置信度最高的一个）
  (6) 过滤置信度低于阈值（0.5）的检测框
  (7) 返回JSON格式结果：[{"label": "cup", "confidence": 0.87, "bbox": [x,y,w,h]}, ...]

后端集成层（Spring Boot YoloClientService）：
- 通过RestTemplate向http://localhost:5000/predict发送POST请求
- 图像先在前端压缩（Canvas重绘至800×800像素，质量80%）再转为Base64编码传输
- 后端接收到Base64后解码为字节数组传给Flask服务
- 解析YOLO返回的结果后执行白名单过滤（仅保留管理员在后台YoloMapping表中配置的有效映射关系）
- 若匹配到唯一分类则直接跳转；若匹配多个则返回候选列表供用户选择；若无匹配则提示"未识别到相关商品"

映射管理的创新设计：
传统做法是将YOLO的80个COCO通用标签硬编码映射到电商分类，这存在两个问题：（1）不同商家的商品品类不同，硬编码缺乏灵活性；（2）COCO标签粒度太粗（如"bottle"无法区分"矿泉水瓶"和"酒瓶"）。本系统通过YoloMappingController提供了动态映射管理界面（admin/yolo-mapping.vue），管理员可以在后台随时增删改"检测标签→系统分类ID"的映射关系，使AI能力能够灵活适配不同的业务场景。'''
ap(doc, yolo_content, 12, ind=0.74)

# 2.5 协同过滤推荐算法
ap(doc, '2.5 协同过滤推荐算法', 13, bw=True, bf=10)
cf_content = '''2.5.1 推荐系统概述与CF算法原理

推荐系统（Recommender System）是利用信息过滤技术预测用户兴趣偏好的软件系统。在电子商务领域，推荐系统对于解决"信息过载"问题、提升用户体验和增加平台转化率具有不可替代的价值。根据Netflix的报告，其推荐系统每年为公司节省超过10亿美元的内容推广成本。

协同过滤（Collaborative Filtering, CF）是推荐系统领域最经典、应用最广泛的算法之一，最早由GroupLens研究组在1994年提出。其核心假设是："如果两个用户在过去喜欢了相同的一些物品，那么他们在未来也很可能喜欢其他相同的物品"（User-based CF）；或者"如果两个物品经常被同一个用户群体一起购买，那么它们之间很可能存在某种关联关系"（Item-based CF）。

2.5.2 Item-based CF在本系统中的实现细节

本系统选择了Item-based CF（基于物品的协同过滤）作为推荐引擎的核心算法，主要原因如下：
（1）物品数量相对稳定且远小于用户数量（通常商品数N << 用户数M），因此物品相似度矩阵的计算和存储开销更低。
（2）物品间的关联关系相对稳定（"购买A的用户有80%也购买了B"这种规律不会频繁变化），可以离线预计算并缓存。
（3）推荐结果易于解释——可以向用户展示"因为您购买了A所以推荐B"的理由，增强用户信任感。

算法实现的五个步骤：

Step 1 - 行为数据采集与加权建模：
系统记录用户的三种交互行为并为每种行为赋予不同的权重分值：
- 浏览行为（点击进入商品详情页）：weight_browse = 1分
- 加购行为（加入购物车）：weight_cart = 3分
- 购买行为（提交订单并完成支付）：weight_buy = 5分

此外引入时间衰减因子 α = e^(-λΔt)，其中λ为衰减系数（本系统设为0.01，约69天半衰期），Δt为行为发生距离当前的天数。最终每个用户对每个物品的交互得分 R[u][i] = Σ(w_k × α_k)，对所有该用户对该物品的历史行为k求和。

Step 2 - 用户-物品交互矩阵构建：
构建 M×N 维的稀疏矩阵R（M为活跃用户数，N为商品总数），其中R[u][i]表示用户u对商品i的综合行为得分。未产生交互的位置记为0。

Step 3 - 物品相似度矩阵计算：
采用修正余弦相似度（Adjusted Cosine Similarity）计算任意两个物品i和j之间的相似度：
Sim(i,j) = Σ(R[u][i] - Ř_u)(R[u][j] - Ř_u) / (√Σ(R[u][i]-Ř_u)² × √Σ(R[u][j]-Ř_u)²)
其中Ř_u 表示用户u对所有物品的平均得分。此公式消除了不同用户打分尺度不一致的影响。

Step 4 - 目标用户推荐列表生成：
对于需要推荐的目标用户u，取出其历史交互过的物品集合 I_u = {i | R[u][i] > 0}。
计算每个未被交互过的物品j的推荐得分：
Score(u,j) = Σ(R[u][i] × Sim(i,j))，对所有 i ∈ I_u 的物品i求和。
取Score值最高的Top-K个（本系统K=20）未交互物品作为推荐结果返回。

Step 5 - 后处理与混合策略：
（1）过滤掉已下架或缺货的商品（status != 1 或 stock <= 0）
（2）与热门榜单结果按7:3比例混合（70%权重来自CF推荐，30%来自销量/好评排序的热门商品），避免推荐结果过于单一或出现"冷门死角"
（3）对新用户（|I_u| == 0）直接返回热门榜单Top-20，绕过CF计算

2.5.3 性能优化措施

由于物品相似度矩阵的计算复杂度为O(N²×M)（N为商品数M为平均交互用户数），当数据量增大时实时计算不可接受。本系统采取以下优化：

（1）离线预计算：通过定时任务（@Scheduled(cron = "0 0 * * * ?")，每小时执行一次）批量更新物品相似度矩阵并存入Redis（key: cf:similarity_matrix），序列化为JSON字符串存储。

（2）在线快速检索：推荐接口触发时从Redis读取预计算的相似度矩阵，反序列化后仅执行Target User的加权求和排序操作（O(K×|I_u|)复杂度），可在50ms内返回结果。

（3）近似最近邻搜索（可选扩展）：若商品数超过10000，可考虑引入ANN（Approximate Nearest Neighbor）库如Faiss或HNSW（Hierarchical Navigable Small World）将物品向量索引化，将相似度搜索从O(N)降低到O(log N)。'''
ap(doc, cf_content, 12, ind=0.74)
doc.add_page_break()

print("Chapter 2 done...")

doc.save(r'd:\shop-vault-project\documents\output\本科毕业论文_完整版_欧阳宇文.docx')
print("[SUCCESS] Part B (Chapter 2) appended!")
